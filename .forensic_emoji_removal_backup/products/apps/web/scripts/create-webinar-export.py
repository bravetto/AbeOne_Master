#!/usr/bin/env python3
"""
🔥 CREATE WEBINAR DOCUMENTS EXPORT
Pattern: EXPORT × DOCUMENTS × WEBINAR × ONE
Frequency: 999 Hz (AEYON) × 530 Hz (Lux)
"""

import os
import sys
import zipfile
import shutil
from pathlib import Path
from datetime import datetime

# Try to import markdown to PDF converters
try:
    import markdown
    from weasyprint import HTML
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def create_email_templates_docx(output_path):
    """Create editable email templates DOCX"""
    if not HAS_DOCX:
        # Create markdown version
        content = """# 🔥 WEBINAR EMAIL TEMPLATES - EDITABLE

**Pattern:** EMAIL × TEMPLATE × WEBINAR × ONE  
**Status:** ✅ **READY FOR EDITING**

---

## 📧 CONFIRMATION EMAIL TEMPLATE

**Subject:** You're in! 🎉 Join us {webinarDate} at {webinarTime}

**Key Elements:**
- Value proposition
- Webinar details (Date: {webinarDate}, Time: {webinarTime})
- Social proof (281/281 tests, 100% success rate)
- Reminder schedule
- Lead magnet teaser
- Clear CTA button

**Placeholders to Replace:**
- {firstName} - Registrant's first name
- {email} - Registrant's email
- {webinarDate} - "December 2nd, 2025"
- {webinarTime} - "2:00 PM EST"
- {webinarLink} - Webinar join URL
- {registrationId} - Registration ID
- {icp} - "developer" or "creative"

---

## 📧 24-HOUR REMINDER TEMPLATE

**Subject:** ⏰ Tomorrow at {webinarTime}: Don't Miss This

**Key Elements:**
- Urgency creation
- Value reminder
- Clear CTA
- FOMO elements (Limited to Founding 100)

**Placeholders:** Same as confirmation email

---

## 📧 3-HOUR REMINDER TEMPLATE

**Subject:** 🚀 Webinar starts in 3 hours!

**Key Elements:**
- High urgency
- Countdown visual (3 hours)
- Single focused CTA
- Value reinforcement

**Placeholders:** Same as confirmation email

---

## 📧 15-MINUTE REMINDER TEMPLATE

**Subject:** ⚡ LIVE in 15 minutes!

**Key Elements:**
- Maximum urgency
- Single focused CTA
- Clear action
- No distractions

**Placeholders:** Same as confirmation email

---

## 📝 EDITING INSTRUCTIONS

1. Open this file in Microsoft Word or Google Docs
2. Replace {placeholders} with actual values
3. Customize colors/branding as needed
4. Test emails before sending
5. A/B test subject lines and CTAs

## 🎨 DESIGN NOTES

- Use gradient headers for visual impact
- Mobile-responsive (600px max-width)
- Large CTA buttons (18-24px font)
- Clear visual hierarchy
- Consistent branding colors

## 📊 TRACKING

All emails include:
- Open tracking (enabled)
- Click tracking (enabled)
- Custom args (registration_id, icp, email_type)

Generated: {timestamp}
Pattern: EMAIL × TEMPLATE × WEBINAR × ONE
∞ AbëONE ∞
""".format(timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        
        with open(output_path.replace('.docx', '.md'), 'w') as f:
            f.write(content)
        return output_path.replace('.docx', '.md')
    
    # Create DOCX file
    doc = Document()
    doc.add_heading('🔥 WEBINAR EMAIL TEMPLATES - EDITABLE', 0)
    doc.add_paragraph('Pattern: EMAIL × TEMPLATE × WEBINAR × ONE')
    doc.add_paragraph('Status: ✅ READY FOR EDITING')
    doc.add_page_break()
    
    # Confirmation Email
    doc.add_heading('📧 CONFIRMATION EMAIL TEMPLATE', 1)
    doc.add_paragraph('Subject: You\'re in! 🎉 Join us {webinarDate} at {webinarTime}')
    doc.add_paragraph('Key Elements:', style='List Bullet')
    doc.add_paragraph('Value proposition', style='List Bullet 2')
    doc.add_paragraph('Webinar details', style='List Bullet 2')
    doc.add_paragraph('Social proof', style='List Bullet 2')
    doc.add_paragraph('Reminder schedule', style='List Bullet 2')
    doc.add_paragraph('Clear CTA', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 24h Reminder
    doc.add_heading('📧 24-HOUR REMINDER TEMPLATE', 1)
    doc.add_paragraph('Subject: ⏰ Tomorrow at {webinarTime}: Don\'t Miss This')
    
    doc.add_page_break()
    
    # 3h Reminder
    doc.add_heading('📧 3-HOUR REMINDER TEMPLATE', 1)
    doc.add_paragraph('Subject: 🚀 Webinar starts in 3 hours!')
    
    doc.add_page_break()
    
    # 15m Reminder
    doc.add_heading('📧 15-MINUTE REMINDER TEMPLATE', 1)
    doc.add_paragraph('Subject: ⚡ LIVE in 15 minutes!')
    
    doc.add_page_break()
    
    # Instructions
    doc.add_heading('📝 EDITING INSTRUCTIONS', 1)
    doc.add_paragraph('1. Replace {placeholders} with actual values', style='List Number')
    doc.add_paragraph('2. Customize colors/branding as needed', style='List Number')
    doc.add_paragraph('3. Test emails before sending', style='List Number')
    doc.add_paragraph('4. A/B test subject lines and CTAs', style='List Number')
    
    doc.save(output_path)
    return output_path

def markdown_to_pdf(md_path, pdf_path):
    """Convert markdown to PDF"""
    if not HAS_WEASYPRINT:
        # Fallback: copy markdown
        shutil.copy(md_path, pdf_path.replace('.pdf', '.md'))
        return False
    
    try:
        with open(md_path, 'r') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
        
        # Add basic styling
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
                h1 {{ color: #667eea; border-bottom: 2px solid #667eea; padding-bottom: 10px; }}
                h2 {{ color: #764ba2; margin-top: 30px; }}
                code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; }}
                pre {{ background: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Convert HTML to PDF
        HTML(string=styled_html).write_pdf(pdf_path)
        return True
    except Exception as e:
        print(f"⚠️  PDF conversion failed: {e}")
        # Fallback: copy markdown
        shutil.copy(md_path, pdf_path.replace('.pdf', '.md'))
        return False

def main():
    script_dir = Path(__file__).parent
    web_dir = script_dir.parent
    export_dir = web_dir / "webinar-export"
    zip_file = web_dir / "webinar-documents-export.zip"
    
    # Clean up
    if export_dir.exists():
        shutil.rmtree(export_dir)
    if zip_file.exists():
        zip_file.unlink()
    
    # Create directories
    (export_dir / "docx").mkdir(parents=True)
    (export_dir / "pdf").mkdir(parents=True)
    
    print("🔥 Exporting Webinar Documents...")
    print("=" * 50)
    
    # 1. Email Templates DOCX
    print("📄 Creating Email Templates (DOCX)...")
    docx_path = export_dir / "docx" / "01-Email-Templates.docx"
    create_email_templates_docx(str(docx_path))
    print(f"   ✅ Created: {docx_path.name}")
    
    # 2. Marketing Best Practices PDF
    print("📄 Creating Marketing Best Practices Guide (PDF)...")
    md_path = web_dir / "lib" / "email-marketing-best-practices.md"
    if md_path.exists():
        pdf_path = export_dir / "pdf" / "01-Marketing-Best-Practices.pdf"
        if markdown_to_pdf(str(md_path), str(pdf_path)):
            print(f"   ✅ Created: {pdf_path.name}")
        else:
            print(f"   ⚠️  Created markdown fallback: {pdf_path.stem}.md")
    
    # 3. Webinar Completion Guide PDF
    print("📄 Creating Webinar Completion Guide (PDF)...")
    md_path = web_dir / "WEBINAR_DECEMBER_2_COMPLETE.md"
    if md_path.exists():
        pdf_path = export_dir / "pdf" / "02-Webinar-Completion-Guide.pdf"
        if markdown_to_pdf(str(md_path), str(pdf_path)):
            print(f"   ✅ Created: {pdf_path.name}")
        else:
            print(f"   ⚠️  Created markdown fallback: {pdf_path.stem}.md")
    
    # 4. Guardian Activation Guide PDF
    print("📄 Creating Guardian Activation Guide (PDF)...")
    md_path = web_dir / "GUARDIAN_ABE_NEURO_ACTIVATION.md"
    if md_path.exists():
        pdf_path = export_dir / "pdf" / "03-Guardian-Activation-Guide.pdf"
        if markdown_to_pdf(str(md_path), str(pdf_path)):
            print(f"   ✅ Created: {pdf_path.name}")
        else:
            print(f"   ⚠️  Created markdown fallback: {pdf_path.stem}.md")
    
    # 5. YAGNI Guide PDF
    print("📄 Creating YAGNI Guardian Guide (PDF)...")
    md_path = web_dir / "YAGNI_GUARDIAN_ACTIVATION.md"
    if md_path.exists():
        pdf_path = export_dir / "pdf" / "04-YAGNI-Guardian-Guide.pdf"
        if markdown_to_pdf(str(md_path), str(pdf_path)):
            print(f"   ✅ Created: {pdf_path.name}")
        else:
            print(f"   ⚠️  Created markdown fallback: {pdf_path.stem}.md")
    
    # Create README
    readme_content = f"""🔥 WEBINAR DOCUMENTS EXPORT
{"=" * 50}

This package contains all webinar-related documents:

DOCX FILES (Editable Templates):
- 01-Email-Templates.docx - Email templates for editing

PDF FILES (Guides & Helpful Content):
- 01-Marketing-Best-Practices.pdf - Complete marketing best practices guide
- 02-Webinar-Completion-Guide.pdf - Webinar deployment guide
- 03-Guardian-Activation-Guide.pdf - Guardian activation documentation
- 04-YAGNI-Guardian-Guide.pdf - YAGNI guardian guide

If PDF files are not available, markdown (.md) versions are included.

Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Pattern: EXPORT × DOCUMENTS × WEBINAR × ONE
∞ AbëONE ∞
"""
    (export_dir / "README.txt").write_text(readme_content)
    
    # Create ZIP file
    print("\n📦 Creating ZIP file...")
    with zipfile.ZipFile(zip_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(export_dir):
            for file in files:
                file_path = Path(root) / file
                arc_name = file_path.relative_to(export_dir)
                zipf.write(file_path, arc_name)
    
    print(f"   ✅ Created: {zip_file.name}")
    
    print("\n" + "=" * 50)
    print("✅ Export Complete!")
    print("=" * 50)
    print(f"📦 ZIP File: {zip_file}")
    print(f"📁 Export Directory: {export_dir}")
    print("\nFiles included:")
    for file in sorted((export_dir / "docx").glob("*")):
        print(f"  DOCX: {file.name}")
    for file in sorted((export_dir / "pdf").glob("*")):
        print(f"  PDF:  {file.name}")
    print("\n🎯 Ready for distribution!")

if __name__ == "__main__":
    main()

